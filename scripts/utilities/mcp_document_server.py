#!/usr/bin/env python3
"""
Enhanced MCP Server with Document Reading Capabilities
Adds PDF, DOCX, and other document processing tools to MCP
"""

import os
import sys
from pathlib import Path
from typing import Optional, Dict, Any
import base64

# Add the project root to Python path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

try:
    from fastmcp import FastMCP
except ImportError:
    print("❌ FastMCP not installed. Install with: pip install fastmcp")
    sys.exit(1)

# Document processing libraries
try:
    import PyPDF2
    from docx import Document
    import fitz  # PyMuPDF for better PDF handling
    from PIL import Image
    import pytesseract  # For OCR if needed
except ImportError as e:
    print(f"❌ Missing document processing libraries: {e}")
    print("Install with: pip install PyPDF2 python-docx PyMuPDF pillow pytesseract")
    sys.exit(1)

# Initialize MCP server
mcp = FastMCP("Kindlemint Document Processor", host="0.0.0.0", port=8012)

@mcp.tool(name="read_pdf", description="Read and extract text from PDF files")
def read_pdf(
    file_path: str,
    extract_images: bool = False,
    page_range: Optional[str] = None
) -> Dict[str, Any]:
    """
    Read PDF file and extract content.
    
    Args:
        file_path: Path to PDF file
        extract_images: Whether to extract images as base64
        page_range: Page range (e.g., "1-5" or "1,3,5")
        
    Returns:
        Dict with extracted content
    """
    try:
        if not os.path.exists(file_path):
            return {"status": "error", "error": f"File not found: {file_path}"}
            
        # Use PyMuPDF for better extraction
        doc = fitz.open(file_path)
        
        # Parse page range
        pages_to_read = []
        if page_range:
            if "-" in page_range:
                start, end = map(int, page_range.split("-"))
                pages_to_read = list(range(start-1, min(end, len(doc))))
            else:
                pages_to_read = [int(p)-1 for p in page_range.split(",")]
        else:
            pages_to_read = list(range(len(doc)))
        
        # Extract content
        content = {
            "text": [],
            "images": [],
            "metadata": {
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "pages": len(doc),
                "file_size_mb": round(os.path.getsize(file_path) / (1024*1024), 2)
            }
        }
        
        for page_num in pages_to_read:
            if page_num >= len(doc):
                continue
                
            page = doc[page_num]
            
            # Extract text
            text = page.get_text()
            content["text"].append({
                "page": page_num + 1,
                "content": text
            })
            
            # Extract images if requested
            if extract_images:
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_data = pix.tobytes("png")
                        img_base64 = base64.b64encode(img_data).decode()
                        content["images"].append({
                            "page": page_num + 1,
                            "index": img_index,
                            "base64": img_base64,
                            "format": "png"
                        })
                    pix = None
        
        doc.close()
        
        return {
            "status": "success",
            "file_path": file_path,
            "content": content,
            "pages_read": len(pages_to_read)
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": str(e),
            "file_path": file_path
        }

@mcp.tool(name="read_docx", description="Read and extract content from DOCX files")
def read_docx(
    file_path: str,
    extract_images: bool = False,
    include_formatting: bool = False
) -> Dict[str, Any]:
    """
    Read DOCX file and extract content.
    
    Args:
        file_path: Path to DOCX file
        extract_images: Whether to extract embedded images
        include_formatting: Whether to include text formatting info
        
    Returns:
        Dict with extracted content
    """
    try:
        if not os.path.exists(file_path):
            return {"status": "error", "error": f"File not found: {file_path}"}
            
        doc = Document(file_path)
        
        content = {
            "paragraphs": [],
            "tables": [],
            "images": [],
            "metadata": {
                "file_size_mb": round(os.path.getsize(file_path) / (1024*1024), 2)
            }
        }
        
        # Extract paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_data = {
                "index": i,
                "text": para.text
            }
            
            if include_formatting:
                para_data["formatting"] = {
                    "style": para.style.name if para.style else None,
                    "alignment": str(para.alignment) if para.alignment else None
                }
                
            content["paragraphs"].append(para_data)
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": []
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data["rows"].append(row_data)
                
            content["tables"].append(table_data)
        
        # Extract images if requested
        if extract_images:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image = rel.target_part.blob
                    img_base64 = base64.b64encode(image).decode()
                    content["images"].append({
                        "base64": img_base64,
                        "content_type": rel.target_part.content_type
                    })
        
        # Extract core properties
        if doc.core_properties:
            content["metadata"].update({
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
            })
        
        return {
            "status": "success",
            "file_path": file_path,
            "content": content
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": str(e),
            "file_path": file_path
        }

@mcp.tool(name="extract_text", description="Extract plain text from various document formats")
def extract_text(
    file_path: str,
    format: Optional[str] = None
) -> Dict[str, Any]:
    """
    Extract plain text from various document formats.
    
    Args:
        file_path: Path to document
        format: Optional format hint (pdf, docx, txt, etc.)
        
    Returns:
        Dict with extracted text
    """
    try:
        if not os.path.exists(file_path):
            return {"status": "error", "error": f"File not found: {file_path}"}
            
        # Detect format if not provided
        if not format:
            ext = Path(file_path).suffix.lower()
            format = ext[1:] if ext else "unknown"
        
        text = ""
        
        if format == "pdf":
            result = read_pdf(file_path)
            if result["status"] == "success":
                text = "\n\n".join([p["content"] for p in result["content"]["text"]])
                
        elif format == "docx":
            result = read_docx(file_path)
            if result["status"] == "success":
                text = "\n\n".join([p["text"] for p in result["content"]["paragraphs"]])
                
        elif format in ["txt", "md", "markdown"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                
        else:
            return {"status": "error", "error": f"Unsupported format: {format}"}
        
        return {
            "status": "success",
            "file_path": file_path,
            "format": format,
            "text": text,
            "length": len(text),
            "lines": text.count("\n") + 1
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": str(e),
            "file_path": file_path
        }

@mcp.tool(name="convert_to_text", description="Convert document to plain text file")
def convert_to_text(
    input_path: str,
    output_path: Optional[str] = None
) -> Dict[str, Any]:
    """
    Convert document to plain text file.
    
    Args:
        input_path: Path to input document
        output_path: Optional output path (defaults to input_path.txt)
        
    Returns:
        Dict with conversion results
    """
    try:
        # Extract text
        result = extract_text(input_path)
        if result["status"] != "success":
            return result
            
        # Determine output path
        if not output_path:
            output_path = str(Path(input_path).with_suffix(".txt"))
        
        # Write text file
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(result["text"])
        
        return {
            "status": "success",
            "input_path": input_path,
            "output_path": output_path,
            "text_length": result["length"],
            "file_size_mb": round(os.path.getsize(output_path) / (1024*1024), 2)
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": str(e),
            "input_path": input_path
        }

@mcp.tool(name="analyze_document", description="Analyze document structure and content")
def analyze_document(
    file_path: str
) -> Dict[str, Any]:
    """
    Analyze document structure, content types, and statistics.
    
    Args:
        file_path: Path to document
        
    Returns:
        Dict with analysis results
    """
    try:
        if not os.path.exists(file_path):
            return {"status": "error", "error": f"File not found: {file_path}"}
            
        ext = Path(file_path).suffix.lower()[1:]
        analysis = {
            "file_info": {
                "path": file_path,
                "format": ext,
                "size_mb": round(os.path.getsize(file_path) / (1024*1024), 2)
            }
        }
        
        if ext == "pdf":
            doc = fitz.open(file_path)
            analysis["pdf_info"] = {
                "pages": len(doc),
                "has_text": any(page.get_text().strip() for page in doc),
                "has_images": any(page.get_images() for page in doc),
                "is_encrypted": doc.is_encrypted,
                "metadata": doc.metadata
            }
            doc.close()
            
        elif ext == "docx":
            doc = Document(file_path)
            analysis["docx_info"] = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
                "has_images": bool(list(doc.part.rels.values()))
            }
            
        # Extract sample content
        text_result = extract_text(file_path, format=ext)
        if text_result["status"] == "success":
            text = text_result["text"]
            analysis["content_stats"] = {
                "total_characters": len(text),
                "total_words": len(text.split()),
                "total_lines": text.count("\n") + 1,
                "sample": text[:500] + "..." if len(text) > 500 else text
            }
        
        return {
            "status": "success",
            "analysis": analysis
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": str(e),
            "file_path": file_path
        }

# Import existing puzzle generation tools
from scripts.utilities.mcp_server import (
    generate_sudoku_book,
    generate_crossword_book,
    create_lead_magnet,
    get_book_stats
)

# Re-register existing tools
mcp.tool(name="generate_sudoku_book", description="Generate a complete Sudoku puzzle book ready for KDP publishing")(generate_sudoku_book)
mcp.tool(name="generate_crossword_book", description="Generate a crossword puzzle book for KDP")(generate_crossword_book)
mcp.tool(name="create_lead_magnet", description="Generate a free puzzle sampler for lead generation")(create_lead_magnet)
mcp.tool(name="get_book_stats", description="Get statistics about generated books")(get_book_stats)

if __name__ == "__main__":
    print("🤖 Starting Enhanced Kindlemint MCP Server...")
    print("📚 Available tools:")
    print("\n📄 Document Processing:")
    print("  - read_pdf: Extract content from PDF files")
    print("  - read_docx: Extract content from DOCX files")
    print("  - extract_text: Get plain text from any document")
    print("  - convert_to_text: Convert documents to text files")
    print("  - analyze_document: Analyze document structure")
    print("\n🎯 Puzzle Generation:")
    print("  - generate_sudoku_book: Create KDP-ready Sudoku books")
    print("  - generate_crossword_book: Create crossword puzzle books")
    print("  - create_lead_magnet: Generate free puzzle samplers")
    print("  - get_book_stats: View generated book statistics")
    print("\n🌐 Server running on http://localhost:8012")
    
    mcp.run()