"""
Convert KPF files to Word DOCX format for KDP publishing
"""
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_kpf_to_docx(kpf_file_path, output_dir="output"):
    """Convert KPF file to proper DOCX format for KDP"""
    
    # Read the KPF file
    with open(kpf_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Set document margins (1 inch all around - standard for KDP)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parse the KPF content
    lines = content.split('\n')
    
    title = ""
    author = ""
    in_content = False
    current_chapter = ""
    
    for line in lines:
        line = line.strip()
        
        if line.startswith("TITLE:"):
            title = line.replace("TITLE:", "").strip()
        elif line.startswith("AUTHOR:"):
            author = line.replace("AUTHOR:", "").strip()
        elif line == "FULL BOOK CONTENT":
            in_content = True
            continue
        elif line.startswith("========"):
            continue
        elif not in_content:
            continue
        
        # Process content
        if line.startswith("CHAPTER") and ":" in line:
            # Chapter heading
            if current_chapter:
                doc.add_page_break()
            
            current_chapter = line
            heading = doc.add_heading(line, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line.startswith("## "):
            # Subheading
            doc.add_heading(line.replace("## ", ""), level=2)
            
        elif line.startswith("*") and line.endswith("*") and len(line) > 10:
            # Italicized text (like riddles)
            p = doc.add_paragraph()
            run = p.add_run(line.replace("*", ""))
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line and not line.startswith("=="):
            # Regular paragraph
            if line:
                doc.add_paragraph(line)
    
    # Add title page at the beginning
    if title:
        # Insert title page at the beginning
        title_para = doc.paragraphs[0]
        title_para.clear()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Inches(0.25)  # Larger font for title
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if author:
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(f"By {author}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    # Save the document
    base_name = os.path.splitext(os.path.basename(kpf_file_path))[0]
    docx_path = os.path.join(output_dir, f"{base_name}.docx")
    
    doc.save(docx_path)
    return docx_path

if __name__ == "__main__":
    kpf_file = "output/kids_puzzle_adventures_the_lost_temple.kpf"
    if os.path.exists(kpf_file):
        docx_file = convert_kpf_to_docx(kpf_file)
        print(f"✓ Converted KPF to DOCX: {docx_file}")
        print(f"✓ File size: {os.path.getsize(docx_file)} bytes")
        print("✓ Ready for KDP upload!")
    else:
        print(f"✗ KPF file not found: {kpf_file}")