"""Professional manuscript formatting utility for KDP publishing."""
import logging
import os
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)

class ManuscriptFormatter:
    """Formats raw text content into professional KDP-ready .docx manuscripts."""
    
    def __init__(self):
        """Initialize the manuscript formatter."""
        logger.info(f"📋 CLASS INIT: ManuscriptFormatter starting initialization")
        self.document = None
        logger.info(f"✅ CLASS INIT COMPLETE: ManuscriptFormatter ready for document creation")
    
    def create_professional_manuscript(
        self,
        book_title: str,
        subtitle: str,
        author: str,
        chapters: List[Dict[str, str]],
        output_path: str,
        include_toc: bool = True,
        **kwargs
    ) -> str:
        """Create a professionally formatted manuscript.
        
        Args:
            book_title: Main title of the book
            subtitle: Subtitle of the book
            author: Author name
            chapters: List of dicts with 'title' and 'content' keys
            output_path: Path where to save the .docx file
            include_toc: Whether to include a table of contents
            **kwargs: Additional formatting options
            
        Returns:
            Path to the created .docx file
        """
        logger.info(f"📋 FUNCTION ENTRY: create_professional_manuscript(title='{book_title}', chapters={len(chapters)})")
        
        try:
            # Create new document
            logger.debug(f"📄 DOCUMENT CREATION: Creating new Word document")
            self.document = Document()
            logger.debug(f"✅ DOCUMENT: New Word document created successfully")
            
            # Set up document margins and formatting
            logger.debug(f"🎨 FORMATTING: Setting up professional document formatting")
            self._setup_document_formatting()
            
            # Add title page
            logger.info(f"📝 CONTENT: Adding title page for '{book_title}' by {author}")
            self._add_title_page(book_title, subtitle, author)
            
            # Add table of contents if requested
            if include_toc:
                logger.info(f"📋 TOC: Adding table of contents with {len(chapters)} chapters")
                self._add_table_of_contents(chapters)
            else:
                logger.debug(f"⏭️ TOC: Skipping table of contents (include_toc=False)")
            
            # Add chapters
            logger.info(f"📚 CHAPTERS: Adding {len(chapters)} chapters to manuscript")
            for i, chapter in enumerate(chapters, 1):
                logger.debug(f"📝 CHAPTER {i}: Adding '{chapter.get('title', 'Untitled')}' ({len(chapter.get('content', ''))} chars)")
                self._add_chapter(chapter['title'], chapter['content'])
            
            # Ensure output directory exists
            output_path = os.path.abspath(output_path)
            output_dir = os.path.dirname(output_path)
            logger.debug(f"📁 DIRECTORY: Ensuring output directory exists: {output_dir}")
            os.makedirs(output_dir, exist_ok=True)
            
            # Save document
            logger.info(f"💾 FILE OPERATION: Saving manuscript to {output_path}")
            self.document.save(output_path)
            
            # Verify file was created
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logger.info(f"✅ SUCCESS: Professional manuscript saved to {output_path} ({file_size} bytes)")
            else:
                logger.error(f"❌ FILE ERROR: Manuscript file was not created at {output_path}")
                raise FileNotFoundError(f"Failed to create manuscript file at {output_path}")
            
            logger.info(f"📤 FUNCTION EXIT: create_professional_manuscript() -> '{output_path}'")
            return output_path
            
        except Exception as e:
            logger.error(f"❌ ERROR: Error creating manuscript: {str(e)}")
            logger.error(f"❌ ERROR DETAILS: Failed during manuscript creation for '{book_title}'")
            logger.info(f"📤 FUNCTION EXIT: create_professional_manuscript() -> Exception")
            raise
    
    def _setup_document_formatting(self):
        """Set up professional document formatting standards."""
        logger.debug(f"📋 FUNCTION ENTRY: _setup_document_formatting()")
        
        # Set document margins (1 inch on all sides)
        logger.debug(f"📏 MARGINS: Setting 1-inch margins on all sides")
        sections = self.document.sections
        margin_count = 0
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            margin_count += 1
        logger.debug(f"✅ MARGINS: Applied margins to {margin_count} document sections")
        
        # Define custom styles
        logger.debug(f"🎨 STYLES: Setting up custom paragraph and text styles")
        styles = self.document.styles
        
        # Chapter title style
        try:
            chapter_style = styles['Chapter Title']
        except KeyError:
            chapter_style = styles.add_style('Chapter Title', WD_STYLE_TYPE.PARAGRAPH)
            chapter_style.font.name = 'Times New Roman'
            chapter_style.font.size = Pt(18)
            chapter_style.font.bold = True
            chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chapter_style.paragraph_format.space_before = Pt(24)
            chapter_style.paragraph_format.space_after = Pt(18)
        
        # Body text style
        try:
            body_style = styles['Body Text Custom']
        except KeyError:
            body_style = styles.add_style('Body Text Custom', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Times New Roman'
            body_style.font.size = Pt(12)
            body_style.paragraph_format.line_spacing = 1.5
            body_style.paragraph_format.space_after = Pt(12)
            body_style.paragraph_format.first_line_indent = Inches(0.5)
    
    def _add_title_page(self, title: str, subtitle: str, author: str):
        """Add a professional title page."""
        # Add vertical spacing
        for _ in range(8):
            self.document.add_paragraph()
        
        # Main title
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        if subtitle:
            subtitle_para = self.document.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.name = 'Times New Roman'
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.italic = True
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        for _ in range(12):
            self.document.add_paragraph()
        
        # Author name
        author_para = self.document.add_paragraph()
        author_run = author_para.add_run(f"By {author}")
        author_run.font.name = 'Times New Roman'
        author_run.font.size = Pt(14)
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break after title page
        self.document.add_page_break()
    
    def _add_table_of_contents(self, chapters: List[Dict[str, str]]):
        """Add a professional table of contents."""
        # TOC title
        toc_title = self.document.add_paragraph()
        toc_title_run = toc_title.add_run("Table of Contents")
        toc_title_run.font.name = 'Times New Roman'
        toc_title_run.font.size = Pt(18)
        toc_title_run.font.bold = True
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.paragraph_format.space_after = Pt(24)
        
        # Introduction entry
        intro_para = self.document.add_paragraph()
        intro_run = intro_para.add_run("Introduction")
        intro_run.font.name = 'Times New Roman'
        intro_run.font.size = Pt(12)
        intro_para.paragraph_format.left_indent = Inches(0.5)
        
        # Add dots and page number
        dots_run = intro_para.add_run(" " + "." * 50 + " ")
        dots_run.font.name = 'Times New Roman'
        dots_run.font.size = Pt(12)
        page_run = intro_para.add_run("3")
        page_run.font.name = 'Times New Roman'
        page_run.font.size = Pt(12)
        
        # Chapter entries
        page_num = 5
        for i, chapter in enumerate(chapters, 1):
            chapter_para = self.document.add_paragraph()
            chapter_run = chapter_para.add_run(f"Chapter {i}: {chapter['title']}")
            chapter_run.font.name = 'Times New Roman'
            chapter_run.font.size = Pt(12)
            chapter_para.paragraph_format.left_indent = Inches(0.5)
            
            # Add dots and page number
            dots_length = max(40 - len(chapter['title']), 10)
            dots_run = chapter_para.add_run(" " + "." * dots_length + " ")
            dots_run.font.name = 'Times New Roman'
            dots_run.font.size = Pt(12)
            page_run = chapter_para.add_run(str(page_num))
            page_run.font.name = 'Times New Roman'
            page_run.font.size = Pt(12)
            
            page_num += 4  # Estimate 4 pages per chapter
        
        # Conclusion entry
        conclusion_para = self.document.add_paragraph()
        conclusion_run = conclusion_para.add_run("Conclusion")
        conclusion_run.font.name = 'Times New Roman'
        conclusion_run.font.size = Pt(12)
        conclusion_para.paragraph_format.left_indent = Inches(0.5)
        
        dots_run = conclusion_para.add_run(" " + "." * 45 + " ")
        dots_run.font.name = 'Times New Roman'
        dots_run.font.size = Pt(12)
        page_run = conclusion_para.add_run(str(page_num))
        page_run.font.name = 'Times New Roman'
        page_run.font.size = Pt(12)
        
        # Page break after TOC
        self.document.add_page_break()
    
    def _add_chapter(self, title: str, content: str):
        """Add a formatted chapter to the document."""
        # Chapter title
        chapter_para = self.document.add_paragraph(title, style='Chapter Title')
        
        # Process chapter content
        self._format_chapter_content(content)
        
        # Page break after each chapter (except the last one)
        self.document.add_page_break()
    
    def _format_chapter_content(self, content: str):
        """Format chapter content with proper paragraphs and styling."""
        # Split content into paragraphs
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        for paragraph_text in paragraphs:
            if not paragraph_text:
                continue
            
            # Check if it's a heading (starts with # or ##)
            if paragraph_text.startswith('##'):
                # Subheading
                heading_text = paragraph_text.replace('##', '').strip()
                heading_para = self.document.add_paragraph()
                heading_run = heading_para.add_run(heading_text)
                heading_run.font.name = 'Times New Roman'
                heading_run.font.size = Pt(14)
                heading_run.font.bold = True
                heading_para.paragraph_format.space_before = Pt(18)
                heading_para.paragraph_format.space_after = Pt(12)
                
            elif paragraph_text.startswith('#'):
                # Main heading (shouldn't happen in chapter content, but handle it)
                heading_text = paragraph_text.replace('#', '').strip()
                heading_para = self.document.add_paragraph()
                heading_run = heading_para.add_run(heading_text)
                heading_run.font.name = 'Times New Roman'
                heading_run.font.size = Pt(16)
                heading_run.font.bold = True
                heading_para.paragraph_format.space_before = Pt(20)
                heading_para.paragraph_format.space_after = Pt(14)
                
            else:
                # Regular paragraph
                para = self.document.add_paragraph()
                
                # Handle bold text (between ** or __)
                parts = self._split_formatted_text(paragraph_text)
                
                for part in parts:
                    if part['bold']:
                        run = para.add_run(part['text'])
                        run.font.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    else:
                        run = para.add_run(part['text'])
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                
                # Apply body text styling
                para.style = 'Body Text Custom'
    
    def _split_formatted_text(self, text: str) -> List[Dict[str, any]]:
        """Split text into parts with formatting information."""
        parts = []
        current_text = ""
        in_bold = False
        i = 0
        
        while i < len(text):
            if i < len(text) - 1 and text[i:i+2] in ['**', '__']:
                # Save current text if any
                if current_text:
                    parts.append({'text': current_text, 'bold': in_bold})
                    current_text = ""
                
                # Toggle bold state
                in_bold = not in_bold
                i += 2
            else:
                current_text += text[i]
                i += 1
        
        # Add remaining text
        if current_text:
            parts.append({'text': current_text, 'bold': in_bold})
        
        return parts if parts else [{'text': text, 'bold': False}]
    
    def format_existing_manuscript(
        self,
        input_path: str,
        output_path: str,
        book_title: str,
        author: str
    ) -> str:
        """Format an existing text manuscript into professional .docx format.
        
        Args:
            input_path: Path to the existing .txt manuscript
            output_path: Path where to save the formatted .docx file
            book_title: Title of the book
            author: Author name
            
        Returns:
            Path to the created .docx file
        """
        logger.info(f"📋 FUNCTION ENTRY: format_existing_manuscript(input='{input_path}', title='{book_title}')")
        
        try:
            # Read the existing manuscript
            logger.info(f"📁 FILE OPERATION: Reading existing manuscript from {input_path}")
            if not os.path.exists(input_path):
                logger.error(f"❌ FILE ERROR: Input manuscript not found at {input_path}")
                raise FileNotFoundError(f"Input manuscript not found: {input_path}")
            
            try:
                with open(input_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                logger.debug(f"📄 FILE CONTENT: Read {len(content)} characters from input manuscript")
            except Exception as e:
                logger.error(f"❌ FILE ERROR: Failed to read input manuscript: {e}")
                raise
            
            # Parse chapters from the content
            logger.debug(f"🔍 PARSING: Extracting chapters from manuscript content")
            chapters = self._parse_chapters_from_text(content)
            logger.info(f"📚 PARSING RESULT: Found {len(chapters)} chapters in manuscript")
            
            # Extract subtitle if present
            subtitle = ""
            if len(chapters) > 0 and 'subtitle' in chapters[0]:
                subtitle = chapters[0]['subtitle']
                logger.debug(f"📝 SUBTITLE: Extracted subtitle '{subtitle}' from first chapter")
            else:
                logger.debug(f"📝 SUBTITLE: No subtitle found in manuscript")
            
            # Create the formatted manuscript
            logger.info(f"🔄 CONVERSION: Converting to professional manuscript format")
            result_path = self.create_professional_manuscript(
                book_title=book_title,
                subtitle=subtitle,
                author=author,
                chapters=chapters,
                output_path=output_path
            )
            
            logger.info(f"✅ SUCCESS: Existing manuscript formatted successfully")
            logger.info(f"📤 FUNCTION EXIT: format_existing_manuscript() -> '{result_path}'")
            return result_path
            
        except Exception as e:
            logger.error(f"❌ ERROR: Error formatting existing manuscript: {str(e)}")
            logger.error(f"❌ ERROR DETAILS: Failed to format manuscript from '{input_path}' to '{output_path}'")
            logger.info(f"📤 FUNCTION EXIT: format_existing_manuscript() -> Exception")
            raise
    
    def _parse_chapters_from_text(self, content: str) -> List[Dict[str, str]]:
        """Parse chapters from raw text content."""
        logger.debug(f"📋 FUNCTION ENTRY: _parse_chapters_from_text({len(content)} chars)")
        
        chapters = []
        lines = content.split('\n')
        current_chapter = None
        current_content = []
        
        logger.debug(f"📄 TEXT ANALYSIS: Processing {len(lines)} lines for chapter detection")
        
        chapter_patterns = ['## Chapter', '# Chapter', 'Chapter']
        chapters_found = 0
        lines_processed = 0
        
        for line_num, line in enumerate(lines, 1):
            line = line.strip()
            lines_processed += 1
            
            # Check if this is a chapter heading
            is_chapter_heading = False
            for pattern in chapter_patterns:
                if line.startswith(pattern) and (pattern != 'Chapter' or ':' in line):
                    is_chapter_heading = True
                    logger.debug(f"📋 CHAPTER DETECTED: Line {line_num} matches pattern '{pattern}': '{line[:50]}...'")
                    break
            
            if is_chapter_heading:
                # Save previous chapter if exists
                if current_chapter:
                    chapter_content = '\n\n'.join(current_content)
                    chapters.append({
                        'title': current_chapter,
                        'content': chapter_content
                    })
                    logger.debug(f"📚 CHAPTER SAVED: '{current_chapter}' with {len(chapter_content)} chars")
                    chapters_found += 1
                
                # Start new chapter
                current_chapter = line.replace('#', '').strip()
                current_content = []
                logger.debug(f"🆕 NEW CHAPTER: Started '{current_chapter}'")
                
            elif current_chapter and line:
                # Add content to current chapter
                current_content.append(line)
        
        # Save last chapter
        if current_chapter:
            chapter_content = '\n\n'.join(current_content)
            chapters.append({
                'title': current_chapter,
                'content': chapter_content
            })
            logger.debug(f"📚 FINAL CHAPTER: Saved '{current_chapter}' with {len(chapter_content)} chars")
            chapters_found += 1
        
        logger.info(f"✅ PARSING COMPLETE: Found {chapters_found} chapters from {lines_processed} lines")
        logger.debug(f"📤 FUNCTION EXIT: _parse_chapters_from_text() -> {len(chapters)} chapters")
        return chapters