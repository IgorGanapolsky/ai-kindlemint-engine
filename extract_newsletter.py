#!/usr/bin/env python3
"""Extract and analyze text from newsletter.docx"""

import os
from docx import Document

def extract_text_from_docx(file_path):
    """Extract all text from a Word document"""
    doc = Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    # Extract text from tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    
    return '\n\n'.join(full_text)

def main():
    file_path = '/home/igorganapolsky/Downloads/newsletter.docx'
    
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return
    
    print("Extracting text from newsletter.docx...\n")
    print("=" * 80)
    
    try:
        text_content = extract_text_from_docx(file_path)
        print(text_content)
        print("\n" + "=" * 80)
        print(f"\nTotal characters extracted: {len(text_content)}")
        
        # Save to text file for easier analysis
        output_file = '/home/igorganapolsky/workspace/git/ai-kindlemint-engine/newsletter_content.txt'
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text_content)
        print(f"\nText content saved to: {output_file}")
        
    except Exception as e:
        print(f"Error reading document: {str(e)}")

if __name__ == "__main__":
    main()